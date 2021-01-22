#coding=utf-8
import re,os,json,urllib.parse,docx,requests,random
from docx.shared import Inches
from docx.oxml.ns import qn
from requests_html import HTMLSession
#爬取知识星球内容
def getTalk(doc,show_comments):
    for i in show_comments:
        tt = upadteText(i['text'])
        try:
            doc.add_paragraph(i['owner']['name']+"➡"+i['repliee']['name']+tt)
        except:
            doc.add_paragraph(i['owner']['name']+":"+tt)
def getData(doc,num):
    with open("./ak2/{}.json".format(num), "r") as f:
        json_dict = json.load(f)
    for i in json_dict['resp_data']['topics']:
        doc.add_paragraph("-----------------")
        # print(i)
        try:
            doc.add_paragraph(upadteText(i['talk']['text']))
        except:
            pass
        try:
            getTalk(doc,i['show_comments'])
        except:
            continue
        try:
            images = i['talk']['images']
            getImages(doc,images)
        except:
            continue
        finally:
            doc.add_paragraph("\n-----------------")
def upadteText(text):
    urlList = re.findall('href="(.*?)"', text)
    tags = re.findall('''<e type="hashtag" hid=".*?" title="(.*?)" />''',text)
    at = re.findall('''<e type="mention" uid=".*?" title="(.*?)" />''',text)
    textUpateUrl = text
    if urlList:
        for i in urlList:
            textUpateUrl = re.sub('''<e type="web" href=".*?" />''',urllib.parse.unquote(i),textUpateUrl,count=1)
    gettags = textUpateUrl
    if tags:
        for i in tags:
            gettags = re.sub('''<e type="hashtag" hid=".*?" title=".*?" />''', urllib.parse.unquote(i), gettags,count=1)
    ats = gettags
    # print(gettags)
    if at:
        for i in at:
            ats = re.sub('''<e type="mention" uid=".*?" title=".*?" />''', urllib.parse.unquote(i), ats,count=1)
    return ats
def getImages(doc,images):#传入doc对象
    pathdir = os.listdir("./images")
    for i in images:
        image = requests.get(i["original"]['url'])
        if image.status_code==200:
            num = random.randint(0, 999)
            abool = True
            while abool:
                if "/"+str(num)+".png" in pathdir:
                    num = random.randint(0, 999)
                    print(num)
                else:
                    abool = False
            with open("./images/{}.png".format(num), "wb") as f:
                f.write(image.content)
                doc.add_picture("./images/{}.png".format(num), width=Inches(3))
                # os.remove("{}.png".format(num))
        else:
            doc.add_paragraph("图片下载失败，地址："+i["original"]['url'])
def getRmTag(text):
    tags = re.findall('''<e type="hashtag" hid=".*?" title="(.*?)" />''',text)
    gettags = re.sub('''<e type="hashtag" hid=".*?" title="(.*?)" />''',"",text)
    if tags:
        return gettags
    else:
        return text
def mkDocx(num):
    global doc
    doc.styles['Normal'].font.name = u'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    getData(doc,num)
def getJson2(group_id,try_num):
    global headers,cookies
    session = HTMLSession()
    x = 0
    y = 0
    q = 1
    try_num2 =try_num
    while x<try_num:
        get_json = session.get("https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20".format(group_id),cookies=cookies, headers=headers)
        json_get = json.loads(get_json.text)
        if get_json.status_code == 200 and json_get['succeeded'] == False:
            x += 1
            print("重复请求次数"+str(x)+"--error_url:"+"https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20".format(group_id))
            if x == 10:
                with open("requests.log","a") as f:
                    f.write("请检查token")
            with open("requests.log", "a") as f:
                f.write("重复请求次数" + str(x) + "--error_url:" + "https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20&".format(group_id)+"\n")

        elif get_json.status_code == 200 and json_get['succeeded'] == True:
            with open("./ak2/1.json", "w", encoding="utf-8") as f:
                json.dump(json_get, f, indent=4)
            x = try_num
            print("https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20  right".format(group_id))
            create_time = json_get['resp_data']['topics'][-1]['create_time']
            while y<try_num2:
                get_json = session.get("https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20&end_time={}".format(group_id,urllib.parse.quote(create_time)),cookies=cookies, headers=headers)
                json_get = json.loads(get_json.text)
                if get_json.status_code == 200 and json_get['succeeded'] == False:
                    y=y+1
                    with open("requests.log","a") as f:
                        f.write("重复请求次数"+str(y)+"--error_url:"+"https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20&end_time={}".format(group_id,urllib.parse.quote(create_time))+"\n")
                    print("重复请求次数"+str(y)+"--error_url:"+"https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20&end_time={}".format(group_id,urllib.parse.quote(create_time)))

                elif get_json.status_code == 200 and json_get['succeeded'] == True:
                    y=0
                    print("https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20&end_time={} right".format(group_id,urllib.parse.quote(create_time)))
                    with open("requests.log","a") as f:
                        f.write("重复请求次数"+str(y)+"--error_url:"+"https://api.zsxq.com/v2/groups/{}/topics?scope=all&count=20&end_time={}\n".format(group_id,urllib.parse.quote(create_time)))

                    create_time = json_get['resp_data']['topics'][-1]['create_time']
                    q +=1
                    with open("./ak2/{}.json".format(str(q)), "w", encoding="utf-8") as f:
                        json.dump(json_get, f, indent=4)
                    if len(json_get['resp_data']['topics']) < 20:
                        y=try_num2
    return q
def getFile(num):
    global cookies
    global headers
    with open("./ak2/{}.json".format(num), "r") as f:
        json_dict = json.load(f)
    for i in json_dict['resp_data']['topics']:
        try:
            for j in i['talk']['files']:
            # print(i['talk']['files'])
                print(j['name'])
                print(j['file_id'])
                # print(j['size'])
                print("a")
                nihao = requests.get('https://api.zsxq.com/v2/files/{}/download_url'.format(j['file_id']),headers=headers,cookies=cookies)
                print(nihao.status_code)
                print("b")
                with open("temp.json", "w", encoding="utf-8") as f:
                    json.dump(nihao, f, indent=4)
        except:
            pass
def rmFile(image=1,file=1):
    if image==1:
        for i in os.listdir("./images"):
            rmI = re.findall("(.*?\.png)",i)
            if rmI:
                os.remove("./images/"+rmI[0])
    if file==1:
        for i in os.listdir("./ak2"):
            rmI = re.findall("(.*?\.json)",i)
            if rmI:
                os.remove("./ak2/"+rmI[0])

if __name__ == '__main__':
    doc = docx.Document()
    headers = {
        'Referer': "https://wx.zsxq.com/",
        'Origin': "https://wx.zsxq.com",
        'User-Agent': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36",#因为是你的cookie登录的，其实headers是什么都无所谓了
        'Host': 'api.zsxq.com'
    }
    cookies = {
        'abtest_env': 'product',
        'zsxq_access_token': 'iamtoken' #token值
    }
    for i in range(getJson2("iamgroupid",10)):#getJson2(https://wx.zsxq.com/dweb2/index/group/后面的数字,重复请求次数)
        mkDocx(i+1)
    rmFile(image=0,file=0)#删除文件 ,,如果需要删除文件，把0改成1
    doc.save(u'结果.docx')

